from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "CCTV-Google-Drive-User-Guide-TH.docx"
LOGO = ROOT / "public" / "images" / "nt-logo.png"

FONT = "Thonburi"
BLUE = "0072CE"
CYAN = "21A8F6"
NAVY = "0B2545"
YELLOW = "FFD200"
INK = "172033"
MUTED = "596579"
LIGHT_BLUE = "EAF6FD"
LIGHT_YELLOW = "FFF8D9"
LIGHT_GRAY = "F3F6F9"
BORDER = "CED8E3"
WHITE = "FFFFFF"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def add_text(doc, text, *, bold=False, color=INK, size=11, after=6, align=None):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after)
    if align is not None:
        paragraph.alignment = align
    set_font(paragraph.add_run(text), size=size, bold=bold, color=color)
    return paragraph


def add_mixed_paragraph(doc, parts, *, after=6):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after)
    for text, bold, color in parts:
        set_font(paragraph.add_run(text), size=11, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    style_name = f"Heading {level}"
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, label, text, tone="blue"):
    colors = {
        "blue": (LIGHT_BLUE, BLUE),
        "yellow": (LIGHT_YELLOW, "9A7500"),
        "gray": (LIGHT_GRAY, MUTED),
    }
    fill, accent = colors[tone]
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=4, after=8, line=1.20)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    set_font(paragraph.add_run(f"{label}  "), size=10.5, bold=True, color=accent)
    set_font(paragraph.add_run(text), size=10.5, color=INK)
    return paragraph


def add_numbering_definition(doc, num_id=41):
    numbering = doc.part.numbering_part.element
    abstract_id = num_id
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_step(doc, num_id, title, detail):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=4, line=1.25)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    set_font(paragraph.add_run(title), size=11, bold=True, color=NAVY)
    if detail:
        set_font(paragraph.add_run(f" — {detail}"), size=11, color=INK)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        set_paragraph_spacing(paragraph, after=0, line=1.12)
        set_font(paragraph.add_run(header), size=9.5, bold=True, color=WHITE)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_tr_pr.append(repeat)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "F8FAFC")
            paragraph = cells[index].paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.15)
            set_font(paragraph.add_run(value), size=9.3, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(paragraph.add_run("หน้า "), size=8.5, color=MUTED)
    run = paragraph.add_run()
    set_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for header in (section.header, section.even_page_header):
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(hp, after=0, line=1.0)
        set_font(hp.add_run("คู่มือผู้ใช้งาน  •  Digital Twin CCTV"), size=8.5, bold=True, color=MUTED)
        p_pr = hp._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), BORDER)
        borders.append(bottom)
        p_pr.append(borders)

    for footer in (section.footer, section.even_page_footer):
        add_page_number(footer.paragraphs[0])


def build_document():
    doc = Document()
    configure_document(doc)
    prepare_num_id = add_numbering_definition(doc, 41)
    share_num_id = add_numbering_definition(doc, 42)
    create_num_id = add_numbering_definition(doc, 43)
    verify_num_id = add_numbering_definition(doc, 44)
    edit_num_id = add_numbering_definition(doc, 45)

    # Cover: editorial_cover pattern, adapted to the Digital Twin brand.
    add_text(doc, "คู่มือการใช้งาน", bold=True, color=BLUE, size=10, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    if LOGO.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_after = Pt(24)
        logo = logo_p.add_run().add_picture(str(LOGO), width=Inches(1.55))
        logo._inline.docPr.set("descr", "National Telecom company logo")
        logo._inline.docPr.set("title", "National Telecom logo")
    add_text(doc, "สร้างกล้อง CCTV", bold=True, color=NAVY, size=28, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "และเชื่อมต่อภาพจาก Google Drive", bold=True, color=BLUE, size=20, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "สำหรับผู้ใช้งานระบบ Digital Twin", color=MUTED, size=12, after=48, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(doc, "สิ่งที่ระบบทำให้", "กล้องที่เชื่อม Google Drive จะดึงภาพล่าสุดจากโฟลเดอร์ ปี/เดือน/วัน และอัปเดตเฉพาะภาพของกล้องทุก 5 วินาที โดยไม่รีเฟรชทั้งหน้า", "blue")
    add_text(doc, "เวอร์ชัน 1.0  |  18 สิงหาคม 2569", color=MUTED, size=9.5, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    add_heading(doc, "1. ก่อนเริ่มใช้งาน", 1)
    add_text(doc, "คู่มือนี้ใช้สำหรับผู้ที่มีสิทธิ์จัดการกล้อง CCTV ในระบบ หากไม่พบปุ่ม “เพิ่มกล้อง” ให้ติดต่อผู้ดูแลระบบเพื่อขอสิทธิ์ cctv.manage")
    add_heading(doc, "สิ่งที่ต้องเตรียม", 2)
    add_step(doc, prepare_num_id, "ข้อมูลกล้อง", "รหัสกล้อง ชื่อ สถานะ อำเภอ และพิกัด (ถ้ามี)")
    add_step(doc, prepare_num_id, "Google Drive root folder", "หนึ่งโฟลเดอร์หลักต่อกล้อง พร้อมโครงสร้าง ปี/เดือน/วัน")
    add_step(doc, prepare_num_id, "สิทธิ์ดูผ่านลิงก์", "ตั้งค่า root folder เป็น Anyone with the link และบทบาท Viewer")
    add_callout(doc, "ผู้ใช้ไม่ต้องมี API Key", "Google Drive API Key เป็นค่าที่ผู้ดูแลระบบตั้งไว้บนเซิร์ฟเวอร์ ผู้ใช้เพียงเตรียมโฟลเดอร์และวาง URL ในฟอร์มกล้อง", "yellow")

    add_heading(doc, "2. เตรียม Google Drive folder", 1)
    add_heading(doc, "2.1 สร้างโครงสร้าง ปี / เดือน / วัน", 2)
    add_text(doc, "สร้างโฟลเดอร์หลักสำหรับกล้อง แล้วสร้างโฟลเดอร์ย่อยด้วยตัวเลขตามลำดับ ตัวอย่างสำหรับกล้อง CCTV-SB-021:")
    tree_lines = [
        "CCTV-SB-021/",
        "└── 2026/",
        "    └── 08/",
        "        └── 18/",
        "            ├── camera_20260818_090000.jpg",
        "            └── camera_20260818_090005.jpg   ← ภาพล่าสุด",
    ]
    for index, line in enumerate(tree_lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        set_paragraph_spacing(p, after=0 if index < len(tree_lines) - 1 else 8, line=1.05)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT_GRAY)
        p_pr.append(shd)
        set_font(p.add_run(line), size=9.5, color=NAVY, name="Arial Unicode MS")
    add_callout(doc, "รูปแบบชื่อโฟลเดอร์", "ปีต้องเป็นตัวเลข 4 หลัก เช่น 2026 หรือ 2569 ส่วนเดือนและวันใช้ 8 หรือ 08 ได้ ระบบจะเรียงวันที่ล่าสุดและเลือกไฟล์ชนิดภาพที่แก้ไขล่าสุด", "gray")

    add_heading(doc, "2.2 แชร์โฟลเดอร์และคัดลอก URL", 2)
    add_step(doc, share_num_id, "คลิกขวาที่ root folder", "เลือก Share หรือแชร์")
    add_step(doc, share_num_id, "เปลี่ยน General access", "เลือก Anyone with the link และกำหนดบทบาท Viewer")
    add_step(doc, share_num_id, "กด Copy link", "URL ที่ได้ควรมีรูปแบบ https://drive.google.com/drive/folders/…")
    add_callout(doc, "สำคัญ", "ใส่ URL ของ root folder ประจำกล้องเท่านั้น ไม่ใช้ URL ของไฟล์ภาพ โฟลเดอร์ปี โฟลเดอร์เดือน หรือโฟลเดอร์วัน", "yellow")

    doc.add_page_break()
    add_heading(doc, "3. สร้างกล้องในระบบ", 1)
    add_step(doc, create_num_id, "เข้าสู่ระบบ Digital Twin", "ใช้บัญชีที่มีสิทธิ์จัดการกล้อง")
    add_step(doc, create_num_id, "เปิดเมนู CCTV", "เข้าสู่หน้าศูนย์ควบคุม CCTV")
    add_step(doc, create_num_id, "กดปุ่ม “เพิ่มกล้อง”", "ฟอร์มเพิ่มกล้องจะแสดงในแผงรายละเอียด")
    add_step(doc, create_num_id, "กรอกข้อมูลกล้อง", "ตรวจสอบรหัสกล้องและชื่อให้ตรงกับหน้างาน")
    add_step(doc, create_num_id, "วาง Google Drive folder URL", "ใช้ URL ของ root folder ที่คัดลอกไว้")
    add_step(doc, create_num_id, "กด “บันทึก”", "ระบบตรวจสอบรูปแบบ URL และสร้างกล้อง")

    add_heading(doc, "ข้อมูลในฟอร์ม", 2)
    add_table(doc,
        ["ช่องข้อมูล", "จำเป็น", "ตัวอย่าง", "คำอธิบาย"],
        [
            ("รหัสกล้อง", "ใช่", "CCTV-SB-021", "ใช้ตัวอักษร ตัวเลข จุด ขีดกลาง หรือขีดล่าง และห้ามซ้ำ"),
            ("ชื่อภาษาไทย", "ใช่", "กล้องหน้าศาลากลาง", "ชื่อที่ผู้ใช้เห็นในรายการกล้อง"),
            ("ชื่อภาษาอังกฤษ", "ไม่", "City Hall Camera", "เว้นว่างได้"),
            ("Google Drive folder URL", "แนะนำ", "https://drive.google.com/drive/folders/…", "root folder ที่มีโครงสร้าง ปี/เดือน/วัน"),
            ("สถานะ", "ใช่", "ออนไลน์", "สถานะการให้บริการของกล้อง"),
            ("อำเภอ", "ไม่", "เมืองสิงห์บุรี", "ช่วยค้นหาและกรองกล้องตามพื้นที่"),
            ("ละติจูด/ลองจิจูด", "ไม่", "14.8912 / 100.4012", "ใช้แสดงตำแหน่งกล้องบนแผนที่"),
        ],
        [1600, 1100, 2600, 4060],
    )
    add_callout(doc, "กล้องเดิมจาก seed data", "กล้องสาธิตเดิมยังคงใช้ภาพเดิมของระบบ และจะไม่ดึงภาพจาก Google Drive จนกว่าจะมีการระบุ Google Drive URL ให้กล้องนั้น", "blue")

    doc.add_page_break()
    add_heading(doc, "4. ตรวจสอบว่าภาพทำงาน", 1)
    add_step(doc, verify_num_id, "ค้นหากล้องที่เพิ่งสร้าง", "ดูจากรหัสหรือชื่อกล้อง")
    add_step(doc, verify_num_id, "ตรวจป้ายใต้ภาพ", "ควรแสดง “Google Drive · 5s”")
    add_step(doc, verify_num_id, "รอการดึงภาพครั้งแรก", "ระบบจะแสดงภาพที่แก้ไขล่าสุดจากวันที่ล่าสุด")
    add_step(doc, verify_num_id, "ทดสอบอัปเดต", "เพิ่มภาพใหม่ในโฟลเดอร์วันปัจจุบัน แล้วรอประมาณ 5 วินาที")
    add_callout(doc, "พฤติกรรมการรีเฟรช", "ระบบรีเฟรชเฉพาะรูปของกล้องที่มี Google Drive URL ไม่โหลดหน้าใหม่ ไม่เปลี่ยนตัวกรอง และไม่รบกวนกล้องอื่น", "blue")

    add_heading(doc, "5. แก้ไขหรือเปลี่ยน URL", 1)
    add_step(doc, edit_num_id, "เลือกกล้องจากรายการ", "เปิดรายละเอียดกล้องทางด้านขวา")
    add_step(doc, edit_num_id, "ไปที่จัดการกล้อง", "กด “แก้ไข”")
    add_step(doc, edit_num_id, "แก้ไข Google Drive folder URL", "วาง URL ใหม่ หรือเว้นว่างเพื่อตัดการเชื่อมต่อ Drive")
    add_step(doc, edit_num_id, "กด “บันทึก”", "ตรวจป้ายแหล่งภาพหลังบันทึก")

    doc.add_page_break()
    add_heading(doc, "6. การแก้ไขปัญหาเบื้องต้น", 1)
    add_table(doc,
        ["อาการ", "แนวทางตรวจสอบ"],
        [
            ("URL ไม่ผ่านการบันทึก", "ตรวจว่าเป็น URL ของ drive.google.com/drive/folders/… และไม่ใช่ลิงก์ไฟล์ภาพ"),
            ("ขึ้นว่าอ่าน Google Drive ไม่ได้", "ตั้ง root folder เป็น Anyone with the link / Viewer แล้วลอง Copy link ใหม่"),
            ("ยังไม่พบภาพในโฟลเดอร์", "ตรวจโครงสร้าง root/ปี/เดือน/วัน และตรวจว่าไฟล์อยู่ในโฟลเดอร์วันโดยตรง"),
            ("ภาพไม่เปลี่ยนหลังอัปโหลด", "รออย่างน้อย 5 วินาที ตรวจว่าอัปโหลดเสร็จ และไฟล์ใหม่มีเวลาแก้ไขล่าสุด"),
            ("ขึ้นว่ายังไม่ได้ตั้งค่า API Key", "ติดต่อผู้ดูแลระบบให้ตั้ง GOOGLE_DRIVE_API_KEY และ restart ระบบ"),
            ("ไม่พบปุ่มเพิ่มหรือแก้ไขกล้อง", "บัญชีไม่มีสิทธิ์ cctv.manage ให้ติดต่อผู้ดูแลระบบ"),
            ("ภาพจากกล้อง seed ยังเป็นภาพเดิม", "เป็นพฤติกรรมที่ถูกต้อง กล้อง seed จะคงภาพเดิมหากไม่ได้ระบุ Google Drive URL"),
        ],
        [2750, 6610],
    )

    add_heading(doc, "7. เช็กลิสต์ก่อนส่งมอบ", 1)
    checklist = [
        "สร้าง root folder แยกสำหรับกล้องแต่ละตัว",
        "สร้างโฟลเดอร์ ปี/เดือน/วัน ด้วยชื่อที่เป็นตัวเลข",
        "อัปโหลดไฟล์ภาพไว้ในโฟลเดอร์วัน",
        "แชร์ root folder เป็น Anyone with the link / Viewer",
        "วาง URL ของ root folder ในฟอร์มกล้อง",
        "เห็นป้าย Google Drive · 5s และภาพล่าสุดบนการ์ดกล้อง",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        set_paragraph_spacing(p, after=5)
        set_font(p.add_run("☐  "), size=11, color=BLUE, name="Arial Unicode MS")
        set_font(p.add_run(item), size=11, color=INK)

    add_callout(doc, "ติดต่อผู้ดูแลระบบเมื่อ", "โฟลเดอร์แชร์ถูกต้องและโครงสร้างครบ แต่ยังพบข้อผิดพลาดเกี่ยวกับ API Key, สิทธิ์ผู้ใช้ หรือระบบไม่สามารถเชื่อมต่อ Google Drive ได้", "yellow")
    add_text(doc, "จบคู่มือ", bold=True, color=NAVY, size=11, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
